#!/usr/bin/env python3
"""Extract text from source documents for requirements evidence passes.

Usage:
    extract_text.py <country> [file ...]

Reads from <country>/sources/, writes to <country>/.extractions/ (scratch,
git-ignored). PDFs get page markers (=== PAGE n ===) so citations can carry
page numbers. XLSX workbooks are dumped sheet-by-sheet as CSV-ish text with
sheet markers. JSON/MD files are copied unchanged.

Options:
    --check    only report extraction quality (pages, chars/page), write nothing
    --ocr      OCR scanned PDFs first (tesseract + spa) and keep the text layer
               (requires: tesseract-ocr, tesseract-ocr-spa, ghostscript, qpdf)
    Damaged PDFs (broken xref/streams) are auto-repaired via qpdf into a temp
    copy — sources are never modified.

Exit code 1 if any requested file fails. Flags scanned/garbled PDFs with a
warning: those need OCR (use --ocr) before evidence work — do not read
through them. `.xls` (xlrd) and `.docx` (python-docx) are supported.
"""

import json
import subprocess
import sys
import tempfile
import unicodedata
from pathlib import Path

REPO = Path(__file__).resolve().parents[2]
WARN_CHARS = {"\ufffd"}  # replacement chars signal bad decode


def ocr_pdf(path: Path) -> Path:
    """Return a path to a text-layer copy of path (OCR with Spanish)."""
    tmp = Path(tempfile.mkdtemp(prefix="ocr-"))
    out = tmp / path.name
    subprocess.run(
        [sys.executable, "-m", "ocrmypdf", "--language", "spa", "--deskew",
         "--force-ocr", "-q", str(path), str(out)],
        check=True,
    )
    return out


def repair_pdf(path: Path) -> Path:
    """Return a path to a qpdf-repaired copy for structurally damaged PDFs.

    Some PDFs render in lenient viewers (pdf.js) but fail strict parsers
    (pypdf) due to a broken xref or malformed streams. qpdf rebuilds the
    cross-reference table; exit code 3 (warnings) still yields a usable file.
    """
    tmp = Path(tempfile.mkdtemp(prefix="repair-"))
    out = tmp / path.name
    subprocess.run(["qpdf", str(path), str(out)], check=False)
    if not out.exists():
        raise RuntimeError("qpdf repair produced no output")
    return out


def out_name(src: Path) -> Path:
    return src.with_suffix(src.suffix + ".txt" if src.suffix != ".txt" else ".txt")


def extract_pdf(path: Path):
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages, 1):
        try:
            text = page.extract_text() or ""
        except Exception as e:  # noqa: BLE001 - record, don't die mid-doc
            text = ""
            print(f"  WARN {path.name} p.{i}: extraction error {e}", file=sys.stderr)
        pages.append(f"=== PAGE {i} ===\n{text}")
    return pages, reader


def pdf_quality(pages) -> str:
    """Heuristic: little or no text => likely scanned; many replacement chars => garbled."""
    n = len(pages)
    if n == 0:
        return "EMPTY"
    chars = sum(len(p) for p in pages) / n
    junk = sum(p.count(c) for p in pages for c in WARN_CHARS)
    if chars < 80:
        return f"SCANNED? ({chars:.0f} chars/page — needs OCR)"
    if junk > 20:
        return f"GARBLED? ({junk} replacement chars)"
    return f"ok ({chars:.0f} chars/page)"


def extract_xlsx(path: Path):
    import openpyxl

    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    blocks = []
    for ws in wb.worksheets:
        lines = [f"=== SHEET {ws.title} ==="]
        for row in ws.iter_rows(values_only=True):
            cells = ["" if c is None else str(c) for c in row]
            if any(c.strip() for c in cells):
                lines.append(";".join(cells))
        blocks.append("\n".join(lines))
    return blocks


def extract_xls(path: Path):
    import xlrd

    wb = xlrd.open_workbook(str(path))
    blocks = []
    for ws in wb.sheets():
        lines = [f"=== SHEET {ws.name} ==="]
        for row in range(ws.nrows):
            cells = ["" if c is None else str(c) for c in ws.row_values(row)]
            if any(c.strip() for c in cells):
                lines.append(";".join(cells))
        blocks.append("\n".join(lines))
    return blocks


def extract_docx(path: Path):
    import docx

    d = docx.Document(str(path))
    parts = ["=== PARAGRAPHS ==="]
    parts += [p.text for p in d.paragraphs if p.text.strip()]
    for i, t in enumerate(d.tables, 1):
        parts.append(f"=== TABLE {i} ===")
        for row in t.rows:
            parts.append(";".join(c.text.strip() for c in row.cells))
    return ["\n".join(parts)]


def main() -> int:
    args = sys.argv[1:]
    check_only = "--check" in args
    use_ocr = "--ocr" in args
    args = [a for a in args if a not in ("--check", "--ocr")]
    if not args:
        print(__doc__)
        return 2
    country, files = args[0], args[1:]
    src_dir = REPO / country / "sources"
    out_dir = REPO / country / ".extractions"
    if not src_dir.is_dir():
        print(f"error: {src_dir} not found", file=sys.stderr)
        return 2

    targets = [src_dir / f for f in files] or sorted(src_dir.iterdir())
    failures = 0
    out_dir.mkdir(exist_ok=True)

    for src in targets:
        if src.name == "README.md" or src.is_dir():
            continue
        try:
            if src.suffix.lower() == ".pdf":
                try:
                    pages, _ = extract_pdf(src)
                except Exception:
                    # Structurally damaged PDF: repair a temp copy via qpdf
                    print(f"{src.name}: pypdf failed -> qpdf repair...", file=sys.stderr)
                    pages, _ = extract_pdf(repair_pdf(src))
                quality = pdf_quality(pages)
                if quality.startswith(("SCANNED", "GARBLED", "EMPTY")) and use_ocr:
                    print(f"{src.name}: scanned -> OCR (spa)...", file=sys.stderr)
                    ocr_path = ocr_pdf(src)
                    pages, _ = extract_pdf(ocr_path)
                    quality = "OCR: " + pdf_quality(pages)
                print(f"{src.name}: {len(pages)} pages, {quality}")
                if quality.startswith(("SCANNED", "GARBLED", "EMPTY")) or (
                    quality.startswith("OCR") and quality[5:].startswith(("SCANNED", "GARBLED", "EMPTY"))
                ):
                    failures += 1
                    continue
                if not check_only:
                    out_dir.joinpath(src.name + ".txt").write_text(
                        "\n\n".join(pages), encoding="utf-8"
                    )
            elif src.suffix.lower() == ".xlsx":
                if not check_only:
                    blocks = extract_xlsx(src)
                    out_dir.joinpath(src.name + ".txt").write_text(
                        "\n\n".join(blocks), encoding="utf-8"
                    )
                print(f"{src.name}: xlsx dumped")
            elif src.suffix.lower() == ".xls":
                if not check_only:
                    blocks = extract_xls(src)
                    out_dir.joinpath(src.name + ".txt").write_text(
                        "\n\n".join(blocks), encoding="utf-8"
                    )
                print(f"{src.name}: xls dumped")
            elif src.suffix.lower() == ".docx":
                if not check_only:
                    blocks = extract_docx(src)
                    out_dir.joinpath(src.name + ".txt").write_text(
                        "\n\n".join(blocks), encoding="utf-8"
                    )
                print(f"{src.name}: docx dumped")
            elif src.suffix.lower() in {".json", ".md", ".txt"}:
                if not check_only:
                    out_dir.joinpath(src.name).write_text(
                        unicodedata.normalize("NFC", src.read_text(encoding="utf-8")),
                        encoding="utf-8",
                    )
                print(f"{src.name}: copied unchanged")
            else:
                print(f"{src.name}: SKIP (unsupported {src.suffix})")
        except Exception as e:  # noqa: BLE001
            failures += 1
            print(f"{src.name}: FAIL {e}", file=sys.stderr)

    return 1 if failures else 0


if __name__ == "__main__":
    sys.exit(main())
